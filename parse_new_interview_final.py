# -*- coding: utf-8 -*-
"""
解析新的面试题Word文档 - 最终版本
简化逻辑，逐行处理
"""
import json
import re
from docx import Document

def parse_interview_docx(file_path):
    """解析面试题Word文档"""
    doc = Document(file_path)

    # 收集所有段落及其索引
    paragraphs = []
    for i in range(len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if text:
            paragraphs.append({'index': i, 'text': text})

    parts = []
    current_part = None
    current_question = None
    question_counter = 0
    part_counter = 0

    # 中文数字标记
    part_cn_nums = ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、']

    # 第一遍：识别所有问题和答案的索引位置
    item_indices = []  # (start_idx, type, content)
    for idx, para in enumerate(paragraphs):
        text = para['text']

        # 检测章节
        if any(text.startswith(cn) for cn in part_cn_nums):
            item_indices.append((idx, 'part', text))
            continue

        # 检测问题
        q_match = re.match(r'^(\d+)\.\s+(.+)$', text)
        if q_match:
            item_indices.append((idx, 'question', q_match.group(2).strip()))
            continue

        # 检测答案标记
        if text.startswith('Answer 英文：'):
            item_indices.append((idx, 'en_answer_start', ''))
            continue
        if text.startswith('Answer 中文：'):
            item_indices.append((idx, 'cn_answer_start', ''))
            continue

    # 按索引排序
    item_indices.sort(key=lambda x: x[0])

    # 第二遍：构建问题结构
    current_q = None
    current_cn_q = None
    en_answer_lines = []
    cn_answer_lines = []

    for item in item_indices:
        idx, itype, content = item

        if itype == 'part':
            # 新章节
            if current_q:
                if current_part:
                    current_part['questions'].append(current_q)
            elif current_part is None:
                current_part = {
                    'partId': f'PART{part_counter}',
                    'partName': content,
                    'questions': []
                }
                part_counter += 1
            current_q = None
            if current_part and current_part['questions']:
                parts.append(current_part)
            current_part = {
                'partId': f'PART{part_counter}',
                'partName': content,
                'questions': []
            }
            part_counter += 1

        elif itype == 'question':
            # 新问题
            if current_q:
                # 保存前一个问题
                if current_cn_q:
                    current_q['questionContentCN'] = current_cn_q
                current_q['sampleAnswer'] = '\n'.join(en_answer_lines)
                current_q['sampleAnswerCN'] = '\n'.join(cn_answer_lines)
                current_part['questions'].append(current_q)

            question_counter += 1
            current_q = {
                'questionNo': f'Q{question_counter}',
                'questionContent': content,
                'questionContentCN': '',
                'sampleAnswer': '',
                'sampleAnswerCN': ''
            }
            current_cn_q = None
            en_answer_lines = []
            cn_answer_lines = []

        elif itype == 'en_answer_start':
            # 开始收集英文答案
            # 找到这个标记和下一个标记之间的所有文本
            next_idx = idx + 1
            while next_idx < len(paragraphs):
                next_text = paragraphs[next_idx]['text']
                # 检查是否遇到中文答案
                if next_text.startswith('Answer 中文：'):
                    break
                # 检查是否遇到新问题
                if re.match(r'^\d+\.\s+', next_text):
                    break
                # 检查是否遇到新章节
                if any(next_text.startswith(cn) for cn in part_cn_nums):
                    break
                # 收集文本
                if next_text and not next_text.startswith('Answer'):
                    en_answer_lines.append(next_text)
                next_idx += 1

        # cn_answer_start不需要处理，因为英文答案的收集已经处理到那里了

    # 保存最后一个问题
    if current_q:
        if current_cn_q is None:
            # 尝试获取中文问题
            pass
        current_q['sampleAnswer'] = '\n'.join(en_answer_lines)
        current_q['sampleAnswerCN'] = '\n'.join(cn_answer_lines)
        if current_part:
            current_part['questions'].append(current_q)

    # 添加最后一个章节
    if current_part and current_part['questions']:
        parts.append(current_part)

    return parts

def main():
    file_path = 'D:/26年FICO模块英文学习网站/面试题 新.docx'

    print("Parsing document...")
    parts = parse_interview_docx(file_path)

    # 统计信息
    total_questions = sum(len(p['questions']) for p in parts)

    result = {
        'documentName': 'SAP FICO Interview Questions - 91 Questions',
        'parts': parts
    }

    output_path = 'D:/26年FICO模块英文学习网站/public/data/interview-qa.json'

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"Done! Parsed {len(parts)} parts with {total_questions} questions")
    print(f"Saved to: {output_path}")

    # 验证特定问题
    print("\n=== Verifying Q34, Q41, Q56, Q57 ===")
    for part in parts:
        for q in part['questions']:
            if q['questionNo'] in ['Q34', 'Q41', 'Q56', 'Q57']:
                en_len = len(q['sampleAnswer'])
                cn_len = len(q['sampleAnswerCN'])
                en_has = 'YES' if en_len > 0 else 'NO'
                cn_has = 'YES' if cn_len > 0 else 'NO'
                print(f"{q['questionNo']}: EN={en_has} ({en_len} chars), CN={cn_has} ({cn_len} chars)")
                if en_len > 0:
                    print(f"  EN: {q['sampleAnswer'][:100]}...")
                if cn_len > 0:
                    print(f"  CN: {q['sampleAnswerCN'][:100]}...")

if __name__ == '__main__':
    main()
