import json
from docx import Document
import re

def parse_word_document(doc_path):
    """解析Word文档，返回按编号排序的题目列表"""
    doc = Document(doc_path)
    questions_dict = {}

    i = 0
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()

        # 检测题目: 数字. 或 数字. 格式 (兼容无空格的情况)
        match = re.match(r'^(\d+)\.\s*(.+)', text)
        if match and len(text) < 500:
            q_num = int(match.group(1))
            question_en = match.group(2).strip()

            # 跳过中文题目
            i += 1
            question_cn = ''
            if i < len(doc.paragraphs):
                para_text = doc.paragraphs[i].text.strip()
                if not para_text.startswith('Answer'):
                    question_cn = para_text
                    i += 1

            # 找英文答案
            en_answer = ''
            if i < len(doc.paragraphs):
                para_text = doc.paragraphs[i].text.strip()
                if 'Answer' in para_text and '英文' in para_text:
                    i += 1
                    while i < len(doc.paragraphs):
                        para_text = doc.paragraphs[i].text.strip()
                        if para_text.startswith('Answer') and '中文' in para_text:
                            break
                        en_answer += para_text + '\n'
                        i += 1

            # 找中文答案
            cn_answer = ''
            if i < len(doc.paragraphs):
                para_text = doc.paragraphs[i].text.strip()
                if 'Answer' in para_text and '中文' in para_text:
                    i += 1
                    while i < len(doc.paragraphs):
                        para_text = doc.paragraphs[i].text.strip()
                        # 检测是否是下一题
                        next_match = re.match(r'^(\d+)\.', para_text)
                        if next_match:
                            break
                        cn_answer += para_text + '\n'
                        i += 1

            questions_dict[q_num] = {
                'questionNo': f'Q{q_num}',
                'questionContent': question_en,
                'questionContentCN': question_cn,
                'sampleAnswer': en_answer.strip(),
                'sampleAnswerCN': cn_answer.strip()
            }
            continue
        else:
            i += 1

    return questions_dict

# 读取当前的JSON文件
with open('面试内容文档_新版.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# 解析Word文档
word_questions = parse_word_document('面试题 新.docx')

print(f'Word文档解析出 {len(word_questions)} 道题目')
print(f'题目编号范围: {min(word_questions.keys())} - {max(word_questions.keys())}')

# 按照Word文档的题目编号更新JSON
# Part结构保持不变，但题目内容从Word文档获取

parts_mapping = [
    ('SCENARIO1', 'General SAP FICO Interview Questions - 通用SAP FICO面试题', 1, 10),
    ('SCENARIO2', 'Automatic Payment Program (APP) Interview Questions - 自动付款程序(APP)面试题', 11, 20),
    ('SCENARIO3', 'Electronic Bank Statement (EBS) Interview Questions - 电子银行对账单(EBS)面试题', 21, 30),
    ('SCENARIO4', 'Dunning Process Interview Questions - 催款过程面试题', 31, 40),
    ('SCENARIO5', 'FI-MM Integration Interview Questions - FI-MM集成面试题', 41, 50),
    ('SCENARIO6', 'FI-SD Integration Interview Questions - FI-SD集成面试题', 51, 60),
    ('SCENARIO7', 'Asset Accounting Interview Questions - 资产核算(AA)面试题', 61, 70),
    ('SCENARIO8', 'Controlling (CO) Interview Questions - 控制模块(CO)面试题', 71, 80),
    ('SCENARIO9', 'Real-Time SAP FICO Scenario-Based Questions - SAP FICO实时场景题', 81, 100),
]

new_parts = []
for part_id, part_name, start_num, end_num in parts_mapping:
    part_questions = []
    for q_num in range(start_num, end_num + 1):
        if q_num in word_questions:
            q = word_questions[q_num].copy()
            # 确保questionNo正确
            q['questionNo'] = f'Q{q_num}'
            part_questions.append(q)
        else:
            print(f'警告: Word文档中缺失Q{q_num}')
    new_parts.append({
        'partId': part_id,
        'partName': part_name,
        'questions': part_questions
    })

# 更新JSON数据
new_data = {
    'documentName': data['documentName'],
    'parts': new_parts
}

# 保存修复后的JSON
with open('面试内容文档_新版.json', 'w', encoding='utf-8') as f:
    json.dump(new_data, f, ensure_ascii=False, indent=2)

print(f'\n已修复并保存到: 面试内容文档_新版.json')

# 显示修复前后的对比
print('\n=== 修复前后对比 (Q56-60) ===')
print('修复前:')
old_all_questions = []
for part in data['parts']:
    old_all_questions.extend(part['questions'])
for i in range(55, min(60, len(old_all_questions))):
    q = old_all_questions[i]
    print(f"  {q['questionNo']}: {q['questionContent'][:60]}...")

print('\n修复后 (基于Word文档):')
for q_num in range(56, 61):
    if q_num in word_questions:
        wq = word_questions[q_num]
        print(f"  Q{q_num}: {wq['questionContent'][:60]}...")
