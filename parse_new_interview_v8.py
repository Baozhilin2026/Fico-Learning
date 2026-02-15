# -*- coding: utf-8 -*-
"""
解析新的面试题Word文档 - 修复版本 v8
修复索引跳过问题
"""
import json
import re
from docx import Document

def parse_interview_docx(file_path):
    """解析面试题Word文档"""
    doc = Document(file_path)

    parts = []
    current_part = None
    current_question = None
    question_counter = 0
    part_counter = 0

    # 中文数字标记
    part_cn_nums = ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、']

    i = 0
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()

        if not text:
            i += 1
            continue

        # 检测章节标题
        is_part_header = any(text.startswith(cn) for cn in part_cn_nums)

        if is_part_header:
            # 保存上一个问题
            if current_question and current_part:
                current_part['questions'].append(current_question)
                current_question = None

            # 保存上一个章节
            if current_part and current_part['questions']:
                parts.append(current_part)

            part_counter += 1
            current_part = {
                'partId': f'PART{part_counter}',
                'partName': text,
                'questions': []
            }
            i += 1
            continue

        # 检测问题 (数字. 英文问题)
        q_match = re.match(r'^(\d+)\.\s+(.+)$', text)

        if q_match:
            # 保存上一个问题
            if current_question and current_part:
                current_part['questions'].append(current_question)

            question_counter += 1
            en_question = q_match.group(2).strip()

            current_question = {
                'questionNo': f'Q{question_counter}',
                'questionContent': en_question,
                'questionContentCN': '',
                'sampleAnswer': '',
                'sampleAnswerCN': ''
            }

            i += 1
            # 读取中文问题 - 不在循环末尾执行i+=1，避免跳过下一行
            while i < len(doc.paragraphs):
                cn_text = doc.paragraphs[i].text.strip()
                if cn_text:
                    if cn_text.startswith('Answer'):
                        # 遇到Answer标记就停止，不增加i（让外层循环处理）
                        break
                    if cn_text.startswith('中文：'):
                        cn_text = cn_text[3:].strip()
                    current_question['questionContentCN'] = cn_text
                i += 1  # 在循环内增加i
            continue  # 继续外层循环，不执行最后的i+=1

        # 检测答案标记 - 使用更精确的检查
        # 优先检查"Answer 中文："因为中文答案通常在后面
        if text.startswith('Answer 中文：') and current_question:
            # 读取中文答案
            i += 1
            answer_lines = []

            while i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()

                # 停止条件
                if re.match(r'^\d+\.\s+', ans_text):
                    break
                if any(ans_text.startswith(cn) for cn in part_cn_nums):
                    break
                if ans_text.startswith('Answer'):
                    break

                if ans_text:
                    answer_lines.append(ans_text)
                i += 1

            current_question['sampleAnswerCN'] = '\n'.join(answer_lines).strip()
            continue

        # 检测英文答案标记
        if text.startswith('Answer 英文：') and current_question:
            # 读取英文答案
            i += 1
            answer_lines = []

            while i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()

                # 停止条件
                if ans_text.startswith('Answer'):
                    break
                if re.match(r'^\d+\.\s+', ans_text):
                    break
                if any(ans_text.startswith(cn) for cn in part_cn_nums):
                    break

                if ans_text:
                    answer_lines.append(ans_text)
                i += 1

            current_question['sampleAnswer'] = '\n'.join(answer_lines).strip()
            continue

        i += 1

    # 保存最后一个问题和章节
    if current_question and current_part:
        current_part['questions'].append(current_question)
    if current_part and current_part['questions']:
        parts.append(current_part)

    return parts

def main():
    file_path = 'D:/26年FICO模块英文学习网站/面试题 新.docx'

    print("Parsing document...")
    parts = parse_interview_docx(file_path)

    # 统计信息
    total_questions = sum(len(p['questions']) for p in parts)

    result = {
        'documentName': 'SAP FICO Interview Questions - 91 Questions',
        'parts': parts
    }

    output_path = 'D:/26年FICO模块英文学习网站/public/data/interview-qa.json'

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"Done! Parsed {len(parts)} parts with {total_questions} questions")
    print(f"Saved to: {output_path}")

    # 验证特定问题
    print("\n=== Verifying Q34, Q41, Q56, Q57 ===")
    for part in parts:
        for q in part['questions']:
            if q['questionNo'] in ['Q34', 'Q41', 'Q56', 'Q57']:
                en_len = len(q['sampleAnswer'])
                cn_len = len(q['sampleAnswerCN'])
                en_has = 'YES' if en_len > 0 else 'NO'
                cn_has = 'YES' if cn_len > 0 else 'NO'
                print(f"{q['questionNo']}: EN={en_has} ({en_len} chars), CN={cn_has} ({cn_len} chars)")
                if en_len > 0:
                    print(f"  EN: {q['sampleAnswer'][:100]}...")
                if cn_len > 0:
                    print(f"  CN: {q['sampleAnswerCN'][:100]}...")

if __name__ == '__main__':
    main()
