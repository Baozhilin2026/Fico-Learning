# -*- coding: utf-8 -*-
"""
手动修复Q34、Q41、Q56、Q57的答案问题
"""
from docx import Document
import json

def fix_specific_answers():
    file_path = 'D:/26年FICO模块英文学习网站/面试题 新.docx'
    json_path = 'D:/26年FICO模块英文学习网站/public/data/interview-qa.json'

    doc = Document(file_path)

    # Load current JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    print("=== Manually extracting answers from Word ===")

    # Q34 English answer (from lines 563-564)
    q34_en = "Goods Receipt (GR) Dr. Inventory A/C, Cr. GR/IR Clearing A/C\nInvoice Receipt (IR)   Dr. GR/IR Clearing A/C, Cr. Vendor A/C"

    # 从Word文档中读取特定区域
    # Q41 Chinese answer should be around lines 657-658
    # Q57 Chinese answer should exist

    # Let's directly scan for Answer patterns and count
    answer_count = 0
    cn_answer_count = 0
    en_answer_count = 0

    for i in range(len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if 'Answer 英文：' in text:
            en_answer_count += 1
        if 'Answer 中文：' in text:
            cn_answer_count += 1

    print(f"Found {en_answer_count} English answer markers, {cn_answer_count} Chinese answer markers")

    # For now, manually add Q34's English answer
    for part in data['parts']:
        for q in part['questions']:
            if q['questionNo'] == 'Q34':
                if not q['sampleAnswer']:
                    q['sampleAnswer'] = q34_en
                    print("Added Q34 English answer")
            elif q['questionNo'] == 'Q41' and not q['sampleAnswerCN']:
                # Q41 Chinese answer: check lines around 656-658
                pass
            elif q['questionNo'] in ['Q56', 'Q57'] and not q['sampleAnswerCN']:
                pass

    # Save
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print("Done! JSON updated")

if __name__ == '__main__':
    fix_specific_answers()
