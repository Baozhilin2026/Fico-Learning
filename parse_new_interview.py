# -*- coding: utf-8 -*-
"""
解析新的面试题Word文档
格式:
- 9个章节，91道题
- 每题包含: 英文问题, 中文翻译, 英文答案, 中文答案
- 答案不使用表格格式
- 适合视力障碍用户的格式
"""
import json
import re
from docx import Document
from docx.oxml.ns import qn

def set_chinese_font(run):
    """设置中文字体"""
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def parse_interview_docx(file_path):
    """解析面试题Word文档"""
    doc = Document(file_path)

    parts = []
    current_part = None
    current_question = None
    question_counter = 0
    part_counter = 0

    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()

        if not text:
            i += 1
            continue

        # 检测章节标题 (中文数字+英文+中文标题)
        part_pattern = r'^([一二三四五六七八九十]+)\.\s*(.+?Questions?)\s*(.+?)$'
        part_match = re.match(part_pattern, text)

        if part_match:
            # 保存上一个问题
            if current_question and current_part:
                current_part['questions'].append(current_question)
                current_question = None

            # 保存上一个章节
            if current_part and current_part['questions']:
                parts.append(current_part)

            part_counter += 1
            cn_num = part_match.group(1)
            en_title = part_match.group(2).strip()
            cn_title = part_match.group(3).strip()

            current_part = {
                'partId': f'PART{part_counter}',
                'partName': f'{cn_num}. {en_title} {cn_title}',
                'questions': []
            }
            i += 1
            continue

        # 检测问题 (数字. 英文问题)
        q_pattern = r'^(\d+)\.\s+(.+)$'
        q_match = re.match(q_pattern, text)

        if q_match:
            # 保存上一个问题
            if current_question and current_part:
                current_part['questions'].append(current_question)

            question_counter += 1
            q_num = q_match.group(1)
            en_question = q_match.group(2).strip()

            current_question = {
                'questionNo': f'Q{question_counter}',
                'questionContent': en_question,
                'questionContentCN': '',
                'sampleAnswer': '',
                'sampleAnswerCN': ''
            }

            i += 1
            # 读取中文问题
            while i < len(doc.paragraphs):
                cn_text = doc.paragraphs[i].text.strip()
                if cn_text and not cn_text.startswith('Answer'):
                    current_question['questionContentCN'] = cn_text
                    i += 1
                    break
                i += 1
            continue

        # 检测答案
        if text.startswith('Answer'):
            # 确定是英文答案还是中文答案
            is_english = '英文' in text or 'English' in text or ('中文' not in text and 'Chinese' not in text)
            is_chinese = '中文' in text or 'Chinese' in text

            if current_question:
                # 跳过Answer行，读取答案内容
                i += 1
                answer_lines = []

                while i < len(doc.paragraphs):
                    ans_text = doc.paragraphs[i].text.strip()

                    # 遇到新问题或章节就停止
                    if re.match(r'^\d+\.\s+', ans_text):
                        break
                    if re.match(r'^[一二三四五六七八九十]+\.\s+', ans_text):
                        break
                    if ans_text.startswith('Answer'):
                        # 遇到另一个Answer，停止
                        break
                    if not ans_text and i < len(doc.paragraphs) - 1:
                        next_text = doc.paragraphs[i+1].text.strip()
                        if next_text.startswith('Answer') or re.match(r'^\d+\.\s+', next_text):
                            i += 1
                            break

                    if ans_text:
                        answer_lines.append(ans_text)
                    i += 1

                answer_text = '\n'.join(answer_lines).strip()

                if is_english or (not is_chinese and not current_question['sampleAnswer']):
                    current_question['sampleAnswer'] = answer_text
                elif is_chinese:
                    current_question['sampleAnswerCN'] = answer_text

            continue

        i += 1

    # 保存最后一个问题和章节
    if current_question and current_part:
        current_part['questions'].append(current_question)
    if current_part and current_part['questions']:
        parts.append(current_part)

    return parts

def clean_answer_text(text):
    """清理答案文本"""
    if not text:
        return text

    # 移除多余空行
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
    # 移除首尾空白
    text = text.strip()

    return text

def main():
    file_path = 'D:/26年FICO模块英文学习网站/面试题 新.docx'

    print("正在解析文档...")
    parts = parse_interview_docx(file_path)

    # 清理答案文本
    for part in parts:
        for q in part['questions']:
            q['sampleAnswer'] = clean_answer_text(q['sampleAnswer'])
            q['sampleAnswerCN'] = clean_answer_text(q['sampleAnswerCN'])

    # 统计信息
    total_questions = sum(len(p['questions']) for p in parts)

    print(f"\n解析完成:")
    print(f"  章节数: {len(parts)}")
    print(f"  总题数: {total_questions}")
    print(f"\n章节列表:")
    for part in parts:
        print(f"  {part['partId']}: {part['partName']} ({len(part['questions'])}题)")

    # 生成JSON
    result = {
        'documentName': 'SAP FICO Interview Questions - 91 Questions',
        'parts': parts
    }

    output_path = 'D:/26年FICO模块英文学习网站/public/data/interview-qa.json'

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"\n已保存到: {output_path}")

    # 显示前几题预览
    print(f"\n前3题预览:")
    for part in parts[:1]:
        for q in part['questions'][:3]:
            print(f"\n{q['questionNo']}: {q['questionContent'][:60]}...")
            print(f"  中文: {q['questionContentCN'][:40]}...")
            print(f"  英文答案长度: {len(q['sampleAnswer'])} 字符")
            print(f"  中文答案长度: {len(q['sampleAnswerCN'])} 字符")

if __name__ == '__main__':
    main()
