# -*- coding: utf-8 -*-
"""
手动修复Q34、Q41、Q56、Q57的答案问题
"""
from docx import Document
import json

def fix_specific_answers():
    file_path = 'D:/26年FICO模块英文学习网站/面试题 新.docx'
    json_path = 'D:/26年FICO模块英文学习网站/public/data/interview-qa.json'

    doc = Document(file_path)

    # Load current JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    print("=== Manually extracting answers from Word ===")

    # Q34: English answer is at lines 563-564
    # Q41: Chinese answer is at lines 657-658
    # Q56: Chinese answer is somewhere after the question
    # Q57: Chinese answer should exist

    # 直接从Word文档中提取这些问题的完整内容
    # 通过查找问题编号来定位
    target_questions = {
        'Q34': 34,
        'Q41': 41,
        'Q56': 56,
        'Q57': 57
    }

    # 存储找到的内容
    found_content = {q: {
        'question': '',
        'en_answer': '',
        'cn_question': '',
        'cn_answer': ''
    } for q in target_questions}

    # 遍历文档
    current_q = None
    i = 0
    while i < len(doc.paragraphs) and current_q is None:
        text = doc.paragraphs[i].text.strip()

        if not text:
            i += 1
            continue

        # 检查是否是目标问题之一
        q_match = None
        for q_no, q_num in target_questions.items():
            if text.startswith(f'{q_num}.'):
                q_match = q_no
                break

        if q_match:
            print(f'Found {q_match} at line {i+1}: {text[:60]}...')
            current_q = q_match
            found_content[current_q]['question'] = text
            i += 1

            # 读取下一行的中文问题
            while i < len(doc.paragraphs):
                cn_text = doc.paragraphs[i].text.strip()
                if cn_text:
                    if cn_text.startswith('中文：'):
                        found_content[current_q]['cn_question'] = cn_text[3:].strip()
                        print(f'  CN question: {cn_text[:50]}')
                        i += 1
                        break
                    # 如果遇到Answer就停止
                    if cn_text.startswith('Answer'):
                        break
                i += 1
            continue

        # 检查是否是答案标记
        if current_q and 'Answer' in text:
            if 'Answer 英文：' in text and not found_content[current_q]['en_answer']:
                print(f'  Found EN answer marker at line {i+1}')
                i += 1
                answer_lines = []
                while i < len(doc.paragraphs):
                    ans_text = doc.paragraphs[i].text.strip()
                    if 'Answer 中文：' in ans_text:
                        found_content[current_q]['en_answer'] = '\n'.join(answer_lines).strip()
                        print(f'  EN answer: {len(answer_lines)} lines')
                        i += 1
                        break
                    if ans_text.startswith('Answer') or ans_text.startswith(f'{list(target_questions.keys())[0]}.'):
                        # 保留这个答案内容，但停止读取（可能是下一个问题）
                        if answer_lines:
                            found_content[current_q]['en_answer'] = '\n'.join(answer_lines).strip()
                            print(f'  EN answer: {len(answer_lines)} lines (ended by next question/answer)')
                        i -= 1  # 回退一行，让外层处理
                        break
                    if ans_text:
                        answer_lines.append(ans_text)
                    i += 1
                continue

            if 'Answer 中文：' in text and not found_content[current_q]['cn_answer']:
                print(f'  Found CN answer marker at line {i+1}')
                i += 1
                answer_lines = []
                while i < len(doc.paragraphs):
                    ans_text = doc.paragraphs[i].text.strip()
                    # 遇到新问题就停止
                    if ans_text.startswith(f'{list(target_questions.keys())[0]}.'):
                        break
                    if ans_text.startswith('Answer'):
                        break
                    if ans_text:
                        answer_lines.append(ans_text)
                    i += 1
                found_content[current_q]['cn_answer'] = '\n'.join(answer_lines).strip()
                print(f'  CN answer: {len(answer_lines)} lines')
                continue

        i += 1

    print('\n=== Extracted content ===')
    for q_no, content in found_content.items():
        print(f'{q_no}:')
        print(f'  EN answer: {len(content[\"en_answer\"])} chars')
        print(f'  CN answer: {len(content[\"cn_answer\"])} chars')

    # 现在更新JSON
    for part in data['parts']:
        for q in part['questions']:
            if q['questionNo'] in found_content:
                content = found_content[q['questionNo']]
                if content['en_answer']:
                    q['sampleAnswer'] = content['en_answer']
                    print(f'Updated {q[\"questionNo\"]} EN answer: {len(content[\"en_answer\"])} chars')
                if content['cn_answer']:
                    q['sampleAnswerCN'] = content['cn_answer']
                    print(f'Updated {q[\"questionNo\"]} CN answer: {len(content[\"cn_answer\"])} chars')

    # 保存
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print('\n=== Updated and saved! ===')

if __name__ == '__main__':
    fix_specific_answers()
