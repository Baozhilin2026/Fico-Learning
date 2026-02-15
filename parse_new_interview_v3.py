# -*- coding: utf-8 -*-
"""
解析新的面试题Word文档
格式:
- 9个章节，91道题
- 每题包含: 英文问题, 中文翻译, 英文答案, 中文答案
- 答案不使用表格格式
- 适合视力障碍用户的格式
"""
import json
import re
from docx import Document

def parse_interview_docx(file_path):
    """解析面试题Word文档"""
    doc = Document(file_path)

    parts = []
    current_part = None
    current_question = None
    question_counter = 0
    part_counter = 0

    # 中文数字标记
    part_cn_nums = ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、']

    i = 0
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()

        if not text:
            i += 1
            continue

        # 检测章节标题
        is_part_header = any(text.startswith(cn) for cn in part_cn_nums)

        if is_part_header:
            # 保存上一个问题
            if current_question and current_part:
                current_part['questions'].append(current_question)
                current_question = None

            # 保存上一个章节
            if current_part and current_part['questions']:
                parts.append(current_part)

            part_counter += 1
            current_part = {
                'partId': f'PART{part_counter}',
                'partName': text,
                'questions': []
            }
            i += 1
            continue

        # 检测问题 (数字. 英文问题)
        q_match = re.match(r'^(\d+)\.\s+(.+)$', text)

        if q_match:
            # 保存上一个问题
            if current_question and current_part:
                current_part['questions'].append(current_question)

            question_counter += 1
            en_question = q_match.group(2).strip()

            current_question = {
                'questionNo': f'Q{question_counter}',
                'questionContent': en_question,
                'questionContentCN': '',
                'sampleAnswer': '',
                'sampleAnswerCN': ''
            }

            i += 1
            # 读取中文问题
            while i < len(doc.paragraphs):
                cn_text = doc.paragraphs[i].text.strip()
                if cn_text and not cn_text.startswith('Answer'):
                    if cn_text.startswith('中文：'):
                        cn_text = cn_text[3:].strip()
                    current_question['questionContentCN'] = cn_text
                    i += 1
                    break
                i += 1
            continue

        # 检测答案的开始
        if text.startswith('Answer 英文') and current_question:
            # 读取英文答案
            i += 1
            answer_lines = []

            while i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()

                # 停止条件
                if ans_text.startswith('Answer 中文'):
                    break
                if re.match(r'^\d+\.\s+', ans_text):
                    break
                if any(ans_text.startswith(cn) for cn in part_cn_nums):
                    break

                if ans_text:
                    answer_lines.append(ans_text)
                i += 1

            current_question['sampleAnswer'] = '\n'.join(answer_lines).strip()
            # 不增加i，因为下一行可能需要处理"Answer 中文"
            continue

        if text.startswith('Answer 中文') and current_question:
            # 读取中文答案
            i += 1
            answer_lines = []

            while i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()

                # 停止条件
                if re.match(r'^\d+\.\s+', ans_text):
                    break
                if any(ans_text.startswith(cn) for cn in part_cn_nums):
                    break

                if ans_text:
                    answer_lines.append(ans_text)
                i += 1

            current_question['sampleAnswerCN'] = '\n'.join(answer_lines).strip()
            continue

        i += 1

    # 保存最后一个问题和章节
    if current_question and current_part:
        current_part['questions'].append(current_question)
    if current_part and current_part['questions']:
        parts.append(current_part)

    return parts

def main():
    file_path = 'D:/26年FICO模块英文学习网站/面试题 新.docx'

    print("Parsing document...")
    parts = parse_interview_docx(file_path)

    # 统计信息
    total_questions = sum(len(p['questions']) for p in parts)

    result = {
        'documentName': 'SAP FICO Interview Questions - 91 Questions',
        'parts': parts
    }

    output_path = 'D:/26年FICO模块英文学习网站/public/data/interview-qa.json'

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"Done! Parsed {len(parts)} parts with {total_questions} questions")
    print(f"Saved to: {output_path}")

    # Preview
    if parts and parts[0]['questions']:
        q = parts[0]['questions'][0]
        print(f"\nPreview Q1:")
        print(f"  Question: {q['questionContent'][:60]}...")
        print(f"  CN: {q['questionContentCN'][:40]}...")
        print(f"  EN Answer: {len(q['sampleAnswer'])} chars")
        print(f"  CN Answer: {len(q['sampleAnswerCN'])} chars")

if __name__ == '__main__':
    main()
