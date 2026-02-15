# -*- coding: utf-8 -*-
"""
手动修复Q41、Q56、Q57的中文答案
"""
from docx import Document
import json

def extract_answer_after_marker(doc, start_line, end_line):
    """提取指定行范围后的答案内容"""
    lines = []
    for i in range(start_line + 1, min(end_line, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if not text:
            continue
        # 停止条件
        if text.startswith('Answer'):
            break
        # 跳过下一个问题
        if text.startswith('57.') or text.startswith('58.') or text.startswith('59.'):
            break
        lines.append(text)
    return '\n'.join(lines)

def fix_answers():
    doc = Document('D:/26年FICO模块英文学习网站/面试题 新.docx')
    json_path = 'D:/26年FICO模块英文学习网站/public/data/interview-qa.json'

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Q41中文答案（657-658行）
    q41_cn = "价格差异会过账到单独的价格差异科目。Standard Price (S): 价格差异单独过账；Moving Average Price (V): 除非价格大幅波动，一般无价格差异。"

    # Q56中文答案（865-870行）
    q56_cn = """折旧范围目的：
账面折旧：按照公司会计政策
税法折旧：为符合政府法规
成本会计折旧：用于内部分析
每个区域可以根据要求过账到不同的分类账/科目。"""

    # Q57中文答案（查找870行之后）
    q57_cn = """资产交易包括：
购置（Acquisition）：通过供应商采购或直接录入（例如：借：资产科目，贷：供应商/银行）
处置（Retirement）：从账簿中移除资产，可能有或无收入（例如：借：累计折旧，贷：资产科目；借：资产收入/银行，贷：资产处置收入）
转移（Transfer）：将资产从一个成本中心或公司代码转移到另一个（例如：内部或公司间资产转移）
通过事务代码处理，如ABZON（购置）、ABAVN（处置）、ABUMN（转移）、ABT1N（公司间转移）"""

    # 更新JSON
    for part in data['parts']:
        for q in part['questions']:
            if q['questionNo'] == 'Q41':
                q['sampleAnswerCN'] = q41_cn
                print("Fixed Q41 CN answer")
            elif q['questionNo'] == 'Q56':
                q['sampleAnswerCN'] = q56_cn
                print("Fixed Q56 CN answer")
            elif q['questionNo'] == 'Q57':
                q['sampleAnswerCN'] = q57_cn
                print("Fixed Q57 CN answer")

    # 保存
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print("\n=== Fixed! ===")
    print("Q41: Added CN answer")
    print("Q56: Added CN answer")
    print("Q57: Added CN answer")
    print("Q34: Already had EN answer added")

if __name__ == '__main__':
    fix_answers()
