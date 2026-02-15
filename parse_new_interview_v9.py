# -*- coding: utf-8 -*-
"""
解析新的面试题Word文档 - 修复版本 v9
完全重写逻辑，使用状态机方式
"""
import json
import re
from docx import Document

def parse_interview_docx(file_path):
    """解析面试题Word文档"""
    doc = Document(file_path)

    parts = []
    current_part = None
    current_question = None
    question_counter = 0
    part_counter = 0

    # 中文数字标记
    part_cn_nums = ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、']

    # 状态机: 'question' -> 'cn_question' -> 'en_answer' -> 'cn_answer' -> 'question'
    state = 'question'

    i = 0
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()

        if not text:
            i += 1
            continue

        # 检测章节标题
        is_part_header = any(text.startswith(cn) for cn in part_cn_nums)
        if is_part_header:
            # 保存上一个问题
            if current_question and current_part:
                current_part['questions'].append(current_question)
                current_question = None

            # 保存上一个章节
            if current_part and current_part['questions']:
                parts.append(current_part)

            part_counter += 1
            current_part = {
                'partId': f'PART{part_counter}',
                'partName': text,
                'questions': []
            }
            state = 'question'
            i += 1
            continue

        # 检测新问题
        q_match = re.match(r'^(\d+)\.\s+(.+)$', text)
        if q_match and state == 'question':
            # 保存上一个问题
            if current_question and current_part:
                current_part['questions'].append(current_question)

            question_counter += 1
            en_question = q_match.group(2).strip()

            current_question = {
                'questionNo': f'Q{question_counter}',
                'questionContent': en_question,
                'questionContentCN': '',
                'sampleAnswer': '',
                'sampleAnswerCN': ''
            }

            state = 'cn_question'
            i += 1
            continue

        # 在cn_question状态下，读取中文问题或检测英文答案
        if state == 'cn_question':
            if text.startswith('Answer 英文：'):
                # 检测到英文答案开始
                state = 'en_answer'
                i += 1
                continue
            elif text.startswith('Answer 中文：'):
                # 没有英文答案，直接到中文答案
                state = 'cn_answer'
                i += 1
                continue
            elif not text.startswith('Answer'):
                # 这是中文问题内容
                if text.startswith('中文：'):
                    text = text[3:].strip()
                current_question['questionContentCN'] = text
                state = 'question'  # 准备读取下一个问题
            i += 1
            continue

        # 在en_answer状态下，收集英文答案直到遇到中文答案
        if state == 'en_answer':
            if text.startswith('Answer 中文：'):
                # 英文答案结束
                state = 'cn_answer'
                i += 1
                continue
            # 收集英文答案内容
            current_question['sampleAnswer'] += ('\n' if current_question['sampleAnswer'] else '') + text
            i += 1
            continue

        # 在cn_answer状态下，收集中文答案
        if state == 'cn_answer':
            # 检查是否需要开始新问题（下一个数字问题）
            if re.match(r'^\d+\.\s+', text):
                state = 'question'
                # 不增加i，让外层循环处理这个新问题
                continue

            # 检查是否遇到下一个英文答案标记
            if text.startswith('Answer 英文：'):
                state = 'en_answer'
                current_question = {
                    'questionNo': f'Q{question_counter}',
                    'questionContent': current_question['questionContent'],
                    'questionContentCN': current_question['questionContentCN'],
                    'sampleAnswer': current_question['sampleAnswer'],
                    'sampleAnswerCN': ''
                }
                question_counter += 1
                i += 1
                continue

            # 检查是否遇到中文答案（跳过）
            if text.startswith('Answer 中文：'):
                i += 1
                continue

            # 收集中文答案内容
            current_question['sampleAnswerCN'] += ('\n' if current_question['sampleAnswerCN'] else '') + text
            i += 1
            continue

        # 默认：增加索引
        i += 1

    # 保存最后一个问题和章节
    if current_question and current_part:
        current_part['questions'].append(current_question)
    if current_part and current_part['questions']:
        parts.append(current_part)

    return parts

def main():
    file_path = 'D:/26年FICO模块英文学习网站/面试题 新.docx'

    print("Parsing document...")
    parts = parse_interview_docx(file_path)

    # 统计信息
    total_questions = sum(len(p['questions']) for p in parts)

    result = {
        'documentName': 'SAP FICO Interview Questions - 91 Questions',
        'parts': parts
    }

    output_path = 'D:/26年FICO模块英文学习网站/public/data/interview-qa.json'

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"Done! Parsed {len(parts)} parts with {total_questions} questions")
    print(f"Saved to: {output_path}")

    # 验证特定问题
    print("\n=== Verifying Q34, Q41, Q56, Q57 ===")
    for part in parts:
        for q in part['questions']:
            if q['questionNo'] in ['Q34', 'Q41', 'Q56', 'Q57']:
                en_len = len(q['sampleAnswer'])
                cn_len = len(q['sampleAnswerCN'])
                en_has = 'YES' if en_len > 0 else 'NO'
                cn_has = 'YES' if cn_len > 0 else 'NO'
                print(f"{q['questionNo']}: EN={en_has} ({en_len} chars), CN={cn_has} ({cn_len} chars)")
                if en_len > 0:
                    print(f"  EN: {q['sampleAnswer'][:100]}...")
                if cn_len > 0:
                    print(f"  CN: {q['sampleAnswerCN'][:100]}...")

if __name__ == '__main__':
    main()
