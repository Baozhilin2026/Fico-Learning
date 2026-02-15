import json
from docx import Document

# 读取JSON文件
with open('面试内容文档_新版.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

all_questions = []
for part in data['parts']:
    all_questions.extend(part['questions'])

# 读取Word文档并正确解析
doc = Document('面试题 新.docx')

questions_word = []
i = 0
while i < len(doc.paragraphs):
    text = doc.paragraphs[i].text.strip()
    # 检测题目: 数字. 格式
    if text and text[0].isdigit() and '. ' in text and len(text) < 300:
        parts = text.split('. ', 1)
        q_num = int(parts[0])
        question_en = parts[1] if len(parts) > 1 else ''

        # 跳过中文题目
        i += 1
        question_cn = ''
        if i < len(doc.paragraphs):
            para_text = doc.paragraphs[i].text.strip()
            # 如果不是Answer开头，应该是中文题目
            if not para_text.startswith('Answer'):
                question_cn = para_text
                i += 1

        # 找英文答案
        en_answer = ''
        if i < len(doc.paragraphs):
            para_text = doc.paragraphs[i].text.strip()
            if 'Answer' in para_text and '英文' in para_text:
                i += 1
                while i < len(doc.paragraphs):
                    para_text = doc.paragraphs[i].text.strip()
                    if para_text.startswith('Answer') and '中文' in para_text:
                        break
                    en_answer += para_text + '\n'
                    i += 1

        # 找中文答案
        cn_answer = ''
        if i < len(doc.paragraphs):
            para_text = doc.paragraphs[i].text.strip()
            if 'Answer' in para_text and '中文' in para_text:
                i += 1
                while i < len(doc.paragraphs):
                    para_text = doc.paragraphs[i].text.strip()
                    # 检测是否是下一题
                    if para_text and para_text[0].isdigit() and '. ' in para_text:
                        break
                    cn_answer += para_text + '\n'
                    i += 1

        questions_word.append({
            'num': q_num,
            'question': question_en,
            'cn_question': question_cn,
            'en_answer': en_answer.strip(),
            'cn_answer': cn_answer.strip()
        })
        continue  # 已经移动到了正确的位置，不需要i+=1
    else:
        i += 1

print(f'Word总题目数: {len(questions_word)}')

# 比较第55-60题 (索引54-59)
print('\n=== Word文档内容 (Q56-60) ===')
for idx in range(55, min(60, len(questions_word))):
    qw = questions_word[idx]
    print(f'Q{idx+1} (Word中是Q{qw["num"]}):')
    print(f'  题目: {qw["question"][:60]}...')
    print(f'  英文答案开头: {qw["en_answer"][:50]}...')
    print(f'  中文答案开头: {qw["cn_answer"][:50]}...')
    print()

print('\n=== JSON文件内容 (Q56-60) ===')
for idx in range(55, min(60, len(all_questions))):
    jq = all_questions[idx]
    print(f'Q{idx+1}:')
    print(f'  题目: {jq["questionContent"][:60]}...')
    print(f'  英文答案开头: {jq["sampleAnswer"][:50]}...')
    print(f'  中文答案开头: {jq["sampleAnswerCN"][:50]}...')
    print()
